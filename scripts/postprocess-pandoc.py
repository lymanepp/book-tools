#!/usr/bin/env python3
"""
postprocess-pandoc.py — Post-build fixup for pandoc-compiled DOCX.

Merges two postprocessors:
  - Scripture book:   inject_headers_footers(), fix_settings(), fix_sections(),
                      fix_chapter_headings(), fix_post_blockquote_style(),
                      fix_copyright_page(), fix_fonts(), fix_title_size().
  - Counseling book:  full table engine (column profiling, width inference,
                      cell shading/borders/margins, header repeat, cantSplit).

Layout produced:
  Title page / TOC              → blank header, no footer
  Chapter opener (1st page)     → blank header, page number footer (centered)
  All other chapter pages       → recto or verso running header, empty footer

  Recto header:  STYLEREF "ChapterTitleOnly"  [tab→right]  PAGE
  Verso header:  PAGE  [tab→right]  Book title

Page geometry (6×9″, mirror margins, twips):
  Page 8640×12960  inner=1080 outer=720 top/bot=1080  Text width=6840

Usage:
    python3 postprocess-pandoc.py INPUT.docx [OUTPUT.docx] --title "Book Title"

    If OUTPUT is omitted, INPUT is overwritten in place.
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

PALATINO   = "Palatino Linotype"
W          = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
TEXT_WIDTH = Twips(6840)   # 6×9″, mirror margins — text block width
TAB_CLEAR  = Twips(4680)   # clears default centre tab stop

TABLE_HEADER_FILL = "404040"   # dark header row
TABLE_RULE_COLOR  = "D0D0D0"   # cell border colour


# ===========================================================================
# Part A — Headers, footers, section wiring
# (from Scripture postprocess-pandoc.py, unchanged)
# ===========================================================================

# ---------------------------------------------------------------------------
# A1: Paragraph / field helpers
# ---------------------------------------------------------------------------

def _noProof_run(child):
    r   = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rpr.append(OxmlElement("w:noProof"))
    r.append(rpr)
    r.append(child)
    return r


def add_field(paragraph, instr: str) -> None:
    """Append a Word field (PAGE, STYLEREF, …) to paragraph."""
    p     = paragraph._p
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    p.append(_noProof_run(begin))

    it = OxmlElement("w:instrText")
    it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    it.text = f" {instr.strip()} "
    p.append(_noProof_run(it))

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    p.append(_noProof_run(sep))
    p.append(_noProof_run(OxmlElement("w:t")))   # empty placeholder

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    p.append(_noProof_run(end))


def add_tab(paragraph) -> None:
    r = OxmlElement("w:r")
    r.append(OxmlElement("w:tab"))
    paragraph._p.append(r)


def add_text(paragraph, text: str) -> None:
    r   = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rpr.append(OxmlElement("w:noProof"))
    r.append(rpr)
    t      = OxmlElement("w:t")
    t.text = text
    r.append(t)
    paragraph._p.append(r)


def set_running_header_tabs(paragraph) -> None:
    paragraph.paragraph_format.first_line_indent = 0
    paragraph.paragraph_format.tab_stops.clear_all()
    paragraph.paragraph_format.tab_stops.add_tab_stop(TAB_CLEAR,  WD_TAB_ALIGNMENT.CLEAR)
    paragraph.paragraph_format.tab_stops.add_tab_stop(TEXT_WIDTH, WD_TAB_ALIGNMENT.RIGHT)


# ---------------------------------------------------------------------------
# A2: Inject headers and footers into section[0]
# ---------------------------------------------------------------------------

def _get_header_paragraph(section, slot: str):
    hf = getattr(section, slot)
    if hf.is_linked_to_previous:
        hf.is_linked_to_previous = False
    return hf.paragraphs[0]


def inject_headers_footers(doc: Document, title: str) -> None:
    """
    Populate all six header/footer slots on section[0].
    fix_sections() propagates the rIds to every subsequent section.
    """
    s0 = doc.sections[0]

    # Recto header: chapter title (via STYLEREF) → page number
    p = _get_header_paragraph(s0, "header")
    p.style = doc.styles["Header"]
    set_running_header_tabs(p)
    add_field(p, 'STYLEREF "ChapterTitleOnly"')
    add_tab(p)
    add_field(p, "PAGE")

    # Verso (even-page) header: page number → book title
    p = _get_header_paragraph(s0, "even_page_header")
    p.style = doc.styles["Header"]
    set_running_header_tabs(p)
    add_field(p, "PAGE")
    add_tab(p)
    add_text(p, title)

    # First-page header: intentionally blank
    p = _get_header_paragraph(s0, "first_page_header")
    p.style = doc.styles["Header"]

    # Default footer: empty
    p = _get_header_paragraph(s0, "footer")
    p.style = doc.styles["Footer"]

    # Even-page footer: empty
    p = _get_header_paragraph(s0, "even_page_footer")
    p.style = doc.styles["Footer"]

    # First-page footer: centered page number (chapter opener)
    p = _get_header_paragraph(s0, "first_page_footer")
    p.style = doc.styles["Footer"]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, "PAGE")

    print(f"  Recto: STYLEREF ChapterTitleOnly [tab] PAGE")
    print(f"  Verso: PAGE [tab] '{title}'")
    print(f"  First-page: blank header, centred PAGE footer")


# ---------------------------------------------------------------------------
# A3: Enable odd/even headers in settings.xml
# ---------------------------------------------------------------------------

def fix_settings(doc: Document) -> None:
    if not doc.settings.odd_and_even_pages_header_footer:
        doc.settings.odd_and_even_pages_header_footer = True
        print("  Enabled odd/even page headers")


# ---------------------------------------------------------------------------
# A4: Wire header/footer rIds across all sectPr blocks
# ---------------------------------------------------------------------------

def fix_sections(doc: Document) -> None:
    """
    Front-matter sections (before first Chapter/Introduction heading):
      blank headers on all slots, empty footers, no titlePg.

    Chapter sections:
      recto/verso/blank headers, page-number/empty footers, titlePg,
      oddPage section break.
    """
    body    = doc.element.body
    s0_refs = _get_hdrftr_rids(doc.sections[0]._sectPr)

    recto   = s0_refs["headerReference", "default"]
    verso   = s0_refs["headerReference", "even"]
    blank   = s0_refs["headerReference", "first"]
    empty   = s0_refs["footerReference", "default"]
    empty_e = s0_refs["footerReference", "even"]
    page    = s0_refs["footerReference", "first"]

    all_sps  = _all_sectPrs(body)
    first_ch = _first_chapter_sectPr(doc, body)
    ch_idx   = all_sps.index(first_ch) if first_ch in all_sps else len(all_sps)

    for i, sp in enumerate(all_sps):
        _strip_hdrftr_refs(sp)
        if i >= ch_idx:
            _ensure(sp, "type", {"w:val": "oddPage"}, first=True)
            _ensure(sp, "titlePg")
            if i == ch_idx:
                _ensure(sp, "pgNumType", {"w:start": "1"})
            _ref(sp, "headerReference", default=recto, even=verso,  first=blank)
            _ref(sp, "footerReference", default=empty,  even=empty_e, first=page)
        else:
            _ref(sp, "headerReference", default=blank, even=blank,  first=blank)
            _ref(sp, "footerReference", default=empty,  even=empty_e, first=empty)

    print(f"  Wired {len(all_sps)} section(s); chapters start at index {ch_idx}")


def _get_hdrftr_rids(sectPr) -> dict:
    rids = {}
    for child in sectPr:
        local = child.tag.split("}")[-1]
        if local in ("headerReference", "footerReference"):
            rids[local, child.get(qn("w:type"))] = child.get(qn("r:id"))
    return rids


def _strip_hdrftr_refs(sectPr) -> None:
    for child in list(sectPr):
        if child.tag.split("}")[-1] in ("headerReference", "footerReference"):
            sectPr.remove(child)


def _ref(sectPr, tag: str, *, default, even, first) -> None:
    for hf_type, rId in (("default", default), ("even", even), ("first", first)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:type"), hf_type)
        el.set(qn("r:id"), rId)
        sectPr.append(el)


def _ensure(sectPr, tag: str, attribs: dict | None = None, first: bool = False):
    el = sectPr.find(f"{{{W}}}{tag}")
    if el is None:
        el = OxmlElement(f"w:{tag}")
        if attribs:
            for k, v in attribs.items():
                el.set(qn(k), v)
        sectPr.insert(0, el) if first else sectPr.append(el)
    return el


def _all_sectPrs(body) -> list:
    sps = [child.find(f".//{{{W}}}sectPr")
           for child in body
           if child.find(f".//{{{W}}}sectPr") is not None]
    doc_sp = body.find(f"{{{W}}}sectPr")
    if doc_sp is not None:
        sps.append(doc_sp)
    return sps


def _first_chapter_sectPr(doc: Document, body):
    """Return the sectPr closing the section containing the first Chapter/Introduction H1."""
    for para in doc.paragraphs:
        pStyle = para._p.find(f".//{{{W}}}pStyle")
        if pStyle is None or not pStyle.get(qn("w:val"), "").startswith("Heading1"):
            continue
        if not para.text.strip().startswith(("Chapter", "Introduction")):
            continue
        found = False
        for child in body:
            if child is para._p:
                found = True
            if found:
                sp = child.find(f".//{{{W}}}sectPr")
                if sp is not None:
                    return sp
        break
    return body.find(f"{{{W}}}sectPr")


# ---------------------------------------------------------------------------
# A5: Split chapter headings for body display vs. running header
# ---------------------------------------------------------------------------
#
# Heading 1 paragraphs with "N.Title" are rewritten so:
#   Body display:    chapter number (large) + line-breaks + title text
#   Running header:  title only   (via STYLEREF "ChapterTitleOnly")
#   TOC entry:       "N. Title"   (TOC \o reads all w:t; w:br has no text)

CHAPTER_RE = re.compile(r"^(\d+)\.(.+)$")


def fix_chapter_headings(doc: Document) -> None:
    STYLE_NAME = "ChapterTitleOnly"
    if STYLE_NAME not in [s.name for s in doc.styles]:
        doc.styles.add_style(STYLE_NAME, WD_STYLE_TYPE.CHARACTER)

    changed = 0
    for para in doc.paragraphs:
        pStyle = para._p.find(f".//{{{W}}}pStyle")
        if pStyle is None or not pStyle.get(qn("w:val"), "").startswith("Heading1"):
            continue
        already = any(
            r.find(f".//{{{W}}}rStyle[@{qn('w:val')}='ChapterTitleOnly']") is not None
            for r in para._p.iter(f"{{{W}}}r")
        )
        if already:
            continue

        full_text = para.text.strip()
        m = CHAPTER_RE.match(full_text)
        if m:
            _rewrite_chapter_heading(para, m.group(1).strip(), m.group(2).strip(), STYLE_NAME)
        else:
            _tag_heading_title(para, full_text, STYLE_NAME)
        changed += 1

    print(f"  Split {changed} chapter heading(s)")


def _rewrite_chapter_heading(para, prefix: str, title: str, style_name: str) -> None:
    p_el = para._p
    pPr  = p_el.find(f"{{{W}}}pPr")
    for child in list(p_el):
        if child is not pPr:
            p_el.remove(child)

    # Large chapter number
    num_run = _plain_run(prefix)
    rpr = OxmlElement("w:rPr")
    for tag, val in (("w:sz", "48"), ("w:szCs", "48")):
        el = OxmlElement(tag)
        el.set(qn("w:val"), val)
        rpr.append(el)
    num_run.insert(0, rpr)
    p_el.append(num_run)

    # White ". " — invisible on page, but TOC \o reads it
    p_el.append(_white_run(". "))

    # Two line breaks (number line + blank line)
    p_el.append(_linebreak_run())
    p_el.append(_linebreak_run())

    # Title run tagged ChapterTitleOnly for STYLEREF
    p_el.append(_styled_run(title, style_name))


def _tag_heading_title(para, full_text: str, style_name: str) -> None:
    p_el = para._p
    pPr  = p_el.find(f"{{{W}}}pPr")
    for child in list(p_el):
        if child is not pPr:
            p_el.remove(child)
    p_el.append(_styled_run(full_text, style_name))


def _plain_run(text: str):
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    return r


def _white_run(text: str):
    r   = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), "FFFFFF")
    rpr.append(clr)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    return r


def _linebreak_run():
    r = OxmlElement("w:r")
    r.append(OxmlElement("w:br"))
    return r


def _styled_run(text: str, style_name: str):
    r      = OxmlElement("w:r")
    rpr    = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), style_name)
    rpr.append(rStyle)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    return r


# ---------------------------------------------------------------------------
# A6: Fix FirstParagraph misapplied after BlockText
# ---------------------------------------------------------------------------

def fix_post_blockquote_style(doc: Document) -> None:
    """
    Pandoc applies FirstParagraph after block quotes.  That is correct after
    a heading, but wrong after a BlockText paragraph — commentary following a
    scripture quote should use BodyText (with first-line indent).
    """
    paras = doc.paragraphs
    count = 0
    for i in range(1, len(paras)):
        prev_el = paras[i - 1]._p.find(f".//{{{W}}}pStyle")
        curr_el = paras[i]._p.find(f".//{{{W}}}pStyle")
        if prev_el is None or curr_el is None:
            continue
        if (prev_el.get(qn("w:val")) == "BlockText" and
                curr_el.get(qn("w:val")) == "FirstParagraph"):
            curr_el.set(qn("w:val"), "BodyText")
            count += 1
    print(f"  Corrected {count} FirstParagraph → BodyText after BlockText")


# ---------------------------------------------------------------------------
# A7: Copyright page — suppress first-line indent, add spacers
# ---------------------------------------------------------------------------

def fix_copyright_page(doc: Document) -> None:
    """
    Between the title-page section break and the TOC heading:
      suppress first-line indent and insert blank BodyText separators
      between copyright items (ISBN pairs stay together).
    """
    body     = doc.element.body
    children = list(body)

    def is_oddpage_para(el):
        return (el.tag == f"{{{W}}}p" and
                el.find(f".//{{{W}}}type[@{qn('w:val')}='oddPage']") is not None)

    def has_toc_heading(el):
        return el.find(f".//{{{W}}}pStyle[@{qn('w:val')}='TOCHeading']") is not None

    start_idx = next((i for i, el in enumerate(children) if is_oddpage_para(el)), None)
    end_idx   = next((i for i, el in enumerate(children) if has_toc_heading(el)), None)

    if start_idx is None or end_idx is None:
        print("  Copyright page anchors not found — skipped")
        return

    region   = [el for el in children[start_idx + 1: end_idx] if el.tag == f"{{{W}}}p"]
    new_els, prev_isbn = [], False

    for p_el in region:
        pStyle     = p_el.find(f".//{{{W}}}pStyle")
        style      = pStyle.get(qn("w:val")) if pStyle is not None else "Normal"
        text       = "".join(t.text or "" for t in p_el.iter(f"{{{W}}}t")).strip()
        has_sectPr = p_el.find(f".//{{{W}}}sectPr") is not None

        if not text and style not in ("BodyText", "FirstParagraph") and not has_sectPr:
            continue

        is_isbn = text.startswith("ISBN:")
        p_el.get_or_add_pPr().first_line_indent = 0

        if new_els and not (is_isbn and prev_isbn) and not has_sectPr:
            new_els.append(_blank_bodytext())

        new_els.append(p_el)
        prev_isbn = is_isbn

    if not new_els:
        print("  Copyright page: nothing to process")
        return

    insert_at = children.index(region[0])
    for p_el in region:
        body.remove(p_el)
    for j, el in enumerate(new_els):
        body.insert(insert_at + j, el)

    spacers = sum(1 for e in new_els if e.find(f"{{{W}}}r") is None)
    print(f"  Copyright page: {len(region)} paragraph(s), {spacers} spacer(s) inserted")


def _blank_bodytext():
    p   = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    ps  = OxmlElement("w:pStyle")
    ps.set(qn("w:val"), "BodyText")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), "0")
    ppr.append(ps)
    ppr.append(ind)
    p.append(ppr)
    return p


# ---------------------------------------------------------------------------
# A8: Normalise theme fonts in body
# ---------------------------------------------------------------------------

def fix_fonts(doc: Document) -> None:
    """Replace theme-font w:rFonts in body with explicit Palatino Linotype."""
    count = 0
    for rFonts in doc.element.body.iter(qn("w:rFonts")):
        if not any("heme" in v for v in rFonts.attrib.values()):
            continue
        for attr in list(rFonts.attrib):
            del rFonts.attrib[attr]
        rFonts.set(qn("w:ascii"),  PALATINO)
        rFonts.set(qn("w:hAnsi"), PALATINO)
        count += 1
    print(f"  Replaced {count} theme-font rFonts" if count else "  No theme-font rFonts found")


# ---------------------------------------------------------------------------
# A9: Title page — 20pt on the first Heading2 (book title)
# ---------------------------------------------------------------------------

def _style_val(paragraph) -> str:
    el = paragraph._p.find(f".//{{{W}}}pStyle")
    return el.get(qn("w:val")) if el is not None else "Normal"


def fix_title_size(doc: Document) -> None:
    if not doc.paragraphs:
        return
    first = doc.paragraphs[0]
    if "Heading2" not in _style_val(first):
        print("  Title size: first paragraph is not Heading2 — skipped")
        return
    for run in first.runs:
        run.font.size = Pt(20)
        rpr = run._r.find(f"{{{W}}}rPr")
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run._r.insert(0, rpr)
        szCs = rpr.find(f"{{{W}}}szCs")
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rpr.append(szCs)
        szCs.set(qn("w:val"), "40")
    print("  Title page: book title set to 20pt")


# ===========================================================================
# Part B — Table engine
# (from counseling postprocess-pandoc-docx.py)
# ===========================================================================

LABEL_WORDS = {
    "element", "phase", "domain", "type", "marker", "pattern", "text", "scope",
    "replacement", "promise", "prohibition", "cognitive", "behavioral",
    "physiological", "generalized", "panic", "social", "health",
    "trauma-related", "clinical", "theological",
}

PROSE_HINT_WORDS = {
    "description", "question", "significance", "practice", "focus",
    "provides", "assigned", "content", "framework", "counterpart", "contributes",
}


@dataclass
class ColumnProfile:
    index:       int
    header:      str
    avg_len:     float
    max_len:     int
    short_ratio: float
    punct_ratio: float
    header_words: int
    labelish:    float
    proseish:    float
    score:       float


# ---------------------------------------------------------------------------
# B1: Cell XML helpers
# ---------------------------------------------------------------------------

def set_cell_margins(cell, top=45, start=80, bottom=45, end=80) -> None:
    tcPr  = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for tag, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        el = tcMar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tcMar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(OxmlElement("w:tblHeader"))


def set_cant_split(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))


def clear_paragraph_borders(paragraph) -> None:
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is not None:
        pPr.remove(pBdr)


def set_cell_shading(cell, fill: str | None) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = tcPr.find(qn("w:shd"))
    if fill is None:
        if shd is not None:
            tcPr.remove(shd)
        return
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)


def set_cell_borders(cell, *, top=None, bottom=None, start=None, end=None) -> None:
    tcPr      = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for side, spec in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        edge = tcBorders.find(qn(f"w:{side}"))
        if spec is None:
            if edge is not None:
                tcBorders.remove(edge)
            continue
        color, size = spec
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            tcBorders.append(edge)
        edge.set(qn("w:val"),   "single")
        edge.set(qn("w:sz"),    str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_table_layout_fixed(table, total_width: float) -> None:
    tblPr  = table._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(Inches(total_width).emu / 635)))


def set_table_cell_spacing(table, dxa: int = 0) -> None:
    tblPr       = table._tbl.tblPr
    cellSpacing = tblPr.find(qn("w:tblCellSpacing"))
    if cellSpacing is None:
        cellSpacing = OxmlElement("w:tblCellSpacing")
        tblPr.append(cellSpacing)
    cellSpacing.set(qn("w:w"),    str(dxa))
    cellSpacing.set(qn("w:type"), "dxa")


# ---------------------------------------------------------------------------
# B2: Column profiling and width inference
# ---------------------------------------------------------------------------

def _normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\n", " ")).strip()


def _text_cells(col_cells: Iterable) -> list[str]:
    return [t for cell in col_cells if (t := _normalize_text(cell.text))]


def _classify_header_words(text: str) -> tuple[float, float]:
    toks = [t.lower() for t in re.findall(r"[A-Za-z0-9']+", text)]
    if not toks:
        return 0.0, 0.0
    return (
        sum(1 for t in toks if t in LABEL_WORDS)     / len(toks),
        sum(1 for t in toks if t in PROSE_HINT_WORDS) / len(toks),
    )


def column_profiles(table) -> list[ColumnProfile]:
    profiles     = []
    ncols        = len(table.columns)
    header_cells = table.rows[0].cells
    data_rows    = table.rows[1:] if len(table.rows) > 1 else table.rows

    for idx in range(ncols):
        header    = _normalize_text(header_cells[idx].text)
        cells     = [row.cells[idx] for row in data_rows if idx < len(row.cells)]
        texts     = _text_cells(cells) or [""]
        lengths   = [len(t) for t in texts]
        avg_len   = sum(lengths) / max(1, len(lengths))
        max_len   = max(lengths) if lengths else 0
        short_ratio = sum(1 for n in lengths if n <= 24) / max(1, len(lengths))
        punct_ratio = sum(t.count(":") + t.count(";") + t.count(",") for t in texts) / max(1.0, sum(lengths) / 50.0)
        header_words = len(re.findall(r"[A-Za-z0-9']+", header))
        labelish, proseish = _classify_header_words(header)
        score = (avg_len * 1.0 + max_len * 0.35 + punct_ratio * 10.0
                 + proseish * 18.0 - short_ratio * 18.0 - labelish * 14.0)
        profiles.append(ColumnProfile(idx, header, avg_len, max_len, short_ratio,
                                      punct_ratio, header_words, labelish, proseish, score))
    return profiles


def _clamp_ratios(raw: list[float], mins: list[float], maxs: list[float]) -> list[float]:
    vals = raw[:]
    for _ in range(10):
        changed = False
        for i, v in enumerate(vals):
            nv = min(max(v, mins[i]), maxs[i])
            if abs(nv - v) > 1e-9:
                vals[i] = nv
                changed = True
        total = sum(vals) or 1
        vals  = [v / total for v in vals]
        if not changed:
            break
    return vals


def infer_widths(table, total_width: float) -> tuple[list[float], str]:
    ncols = len(table.columns)
    if ncols == 1:
        return [total_width], "left"

    profiles = column_profiles(table)

    def prose_weight(p: ColumnProfile) -> float:
        return (p.avg_len * 1.0 + p.max_len * 0.30 + p.punct_ratio * 10.0
                + p.proseish * 18.0 - p.short_ratio * 10.0 - p.labelish * 10.0)

    first_align = "center" if (profiles[0].short_ratio >= 0.7 or profiles[0].labelish > 0.2) else "left"
    weights = [max(8.0, prose_weight(p)) for p in profiles]
    total   = sum(weights)
    ratios  = [w / total for w in weights]

    if ncols == 2:
        if profiles[0].short_ratio > 0.65 and profiles[1].avg_len > profiles[0].avg_len * 1.4:
            ratios = [0.34, 0.66]
        ratios = _clamp_ratios(ratios, [0.28, 0.42], [0.46, 0.72])

    elif ncols == 3:
        first_label = profiles[0].short_ratio > 0.85 and profiles[0].avg_len < 18
        if first_label:
            col1 = 0.20
            w2   = max(12.0, prose_weight(profiles[1]))
            w3   = max(12.0, prose_weight(profiles[2]))
            pt   = w2 + w3
            rem  = 1.0 - col1
            r2   = min(max(rem * w2 / pt, 0.26), 0.38)
            r3   = min(max(rem * w3 / pt, 0.42), 0.54)
            scale = rem / (r2 + r3)
            ratios = [col1, r2 * scale, r3 * scale]
        else:
            if prose_weight(profiles[2]) > prose_weight(profiles[1]) * 1.15:
                ratios = [0.18, 0.30, 0.52]
            ratios = _clamp_ratios(ratios, [0.16, 0.22, 0.28], [0.24, 0.38, 0.56])

        widths = [round(total_width * r, 2) for r in ratios]
        if first_label and widths[0] < 0.95:
            delta    = 0.95 - widths[0]
            widths[0] = 0.95
            pw2      = max(12.0, prose_weight(profiles[1]))
            pw3      = max(12.0, prose_weight(profiles[2]))
            pt2      = pw2 + pw3
            widths[1] = round(widths[1] - delta * pw2 / pt2, 2)
            widths[2] = round(widths[2] - delta * pw3 / pt2, 2)
        widths[-1] = round(total_width - sum(widths[:-1]), 2)
        return widths, first_align

    elif ncols == 4:
        first_label = profiles[0].short_ratio > 0.85 and profiles[0].avg_len < 18
        if first_label:
            col1   = 0.16
            pws    = [max(10.0, prose_weight(p)) for p in profiles[1:]]
            pt     = sum(pws)
            rem    = 1.0 - col1
            prs    = _clamp_ratios([rem * w / pt for w in pws],
                                   [0.18, 0.20, 0.24], [0.28, 0.30, 0.40])
            scale  = rem / sum(prs)
            ratios = [col1] + [r * scale for r in prs]
        else:
            ratios = _clamp_ratios(ratios,
                                   [0.10, 0.16, 0.18, 0.22],
                                   [0.22, 0.28, 0.30, 0.42])
    else:
        base   = 1.0 / ncols
        ratios = _clamp_ratios(ratios, [base * 0.6] * ncols, [base * 1.7] * ncols)

    widths      = [round(total_width * r, 2) for r in ratios]
    widths[-1]  = round(total_width - sum(widths[:-1]), 2)
    return widths, first_align


# ---------------------------------------------------------------------------
# B3: Apply layout to a single table
# ---------------------------------------------------------------------------

def apply_table_layout(table, widths: list[float], total_width: float,
                        first_col_align: str = "left") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit   = False
    set_table_layout_fixed(table, total_width)
    set_table_cell_spacing(table, 0)

    for row in table.rows:
        set_cant_split(row)
    if table.rows:
        set_repeat_table_header(table.rows[0])

    for r_idx, row in enumerate(table.rows):
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            if idx == 0 and len(widths) in (2, 3, 4):
                set_cell_margins(cell, top=42, start=70, bottom=42, end=70)
            else:
                set_cell_margins(cell, top=42, start=82, bottom=42, end=82)

            top_spec   = (TABLE_RULE_COLOR, 6) if r_idx == 0 else None
            start_spec = (TABLE_RULE_COLOR, 6) if idx == 0   else None
            set_cell_shading(cell, TABLE_HEADER_FILL if r_idx == 0 else None)
            set_cell_borders(cell,
                             top=top_spec,
                             bottom=(TABLE_RULE_COLOR, 6),
                             start=start_spec,
                             end=(TABLE_RULE_COLOR, 6))

            for p in cell.paragraphs:
                try:
                    p.style = "Table Header Text" if r_idx == 0 else "Table Text"
                except Exception:
                    pass
                clear_paragraph_borders(p)
                p.paragraph_format.keep_together = True
                if r_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                                   if idx == 0 and first_col_align == "center"
                                   else WD_ALIGN_PARAGRAPH.LEFT)

    for idx, width in enumerate(widths):
        if idx < len(table.columns):
            table.columns[idx].width = Inches(width)


# ---------------------------------------------------------------------------
# B4: Table-level pagination helpers
# ---------------------------------------------------------------------------

def keep_table_with_intro(doc: Document) -> None:
    tables = {tbl._tbl for tbl in doc.tables}
    for p in doc.paragraphs:
        nxt = p._p.getnext()
        if nxt is not None and nxt in tables and p.text.strip():
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_after    = Pt(4)


def add_spacing_after_tables(doc: Document) -> None:
    tables = {tbl._tbl for tbl in doc.tables}
    for p in doc.paragraphs:
        prev = p._p.getprevious()
        if prev is not None and prev in tables:
            p.paragraph_format.space_before = Pt(8)


def keep_entire_table_together(table) -> None:
    last_row_idx = len(table.rows) - 1
    for r_idx, row in enumerate(table.rows):
        is_last = r_idx == last_row_idx
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.keep_together  = True
                p.paragraph_format.keep_with_next = not is_last


# ---------------------------------------------------------------------------
# B5: Document-level table pass
# ---------------------------------------------------------------------------

def _usable_text_width(section) -> float:
    width  = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
    gutter = getattr(section, "gutter", None)
    if gutter is not None:
        try:
            width -= gutter.inches
        except Exception:
            pass
    return max(width, 1.0)


def process_tables(doc: Document) -> None:
    if not doc.tables:
        print("  No tables found — skipping table pass")
        return
    total_width = _usable_text_width(doc.sections[0]) if doc.sections else 4.5
    for i, table in enumerate(doc.tables):
        widths, first_align = infer_widths(table, total_width)
        apply_table_layout(table, widths, total_width, first_align)
        keep_entire_table_together(table)
    keep_table_with_intro(doc)
    add_spacing_after_tables(doc)
    print(f"  Processed {len(doc.tables)} table(s)")


# ===========================================================================
# Main
# ===========================================================================

def main() -> None:
    ap = argparse.ArgumentParser(
        description="Post-build fixup for pandoc-compiled DOCX."
    )
    ap.add_argument("input",
                    help="Input DOCX (pandoc output)")
    ap.add_argument("output", nargs="?",
                    help="Output DOCX (defaults to overwriting input)")
    ap.add_argument("--title", required=True,
                    help="Book title for the verso running header")
    args = ap.parse_args()

    doc    = Document(args.input)
    output = args.output or args.input

    print("Processing tables...")
    process_tables(doc)

    print("Building headers and footers...")
    inject_headers_footers(doc, args.title)
    fix_settings(doc)

    print("Wiring sections...")
    fix_sections(doc)

    print("Splitting chapter headings...")
    fix_chapter_headings(doc)

    print("Fixing post-blockquote paragraph styles...")
    fix_post_blockquote_style(doc)

    print("Fixing copyright page...")
    fix_copyright_page(doc)

    print("Normalising fonts...")
    fix_fonts(doc)

    print("Fixing title page size...")
    fix_title_size(doc)

    doc.save(output)
    print(f"Saved → {output}")


if __name__ == "__main__":
    main()
