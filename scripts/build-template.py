#!/usr/bin/env python3
"""
build-template.py — Build the pandoc reference-template.docx for a book.

All meaningful typographic choices are driven by arguments supplied from
book.env via docx.sh.  The script has no hardcoded book-specific values;
every default is a deliberate starting point that any book can override.

Usage (called automatically by docx.sh):
    python3 build-template.py [output-path] [options]

Options correspond 1-to-1 with book.env variables.  See docx.sh for how
the mapping is done.  You can also call the script directly for testing:

    python3 build-template.py reference-template.docx \\
        --body-font "Palatino Linotype" \\
        --body-size 11 \\
        --quote-preset scripture

Typography variables (all optional — defaults shown):
    --body-font         FONT    Body text font family.   [Palatino Linotype]
    --body-size         PT      Body text size in pt.    [11]
    --body-align        ALIGN   justify | left           [justify]
    --first-line-indent IN      First-line indent in in. [0.25]
    --line-spacing      MULT    Line spacing multiplier. [1.0]
    --h1-size           PT      Heading 1 size.          [16]
    --h2-size           PT      Heading 2 size.          [13]
    --h3-size           PT      Heading 3 size.          [12]
    --footnote-size     PT      Footnote text size.      [9]
    --table-body-font   FONT    Table body cell font.    [Georgia]
    --table-body-size   PT      Table body cell size.    [9.9]
    --table-header-font FONT    Table header font.       [Arial]
    --table-header-size PT      Table header size.       [10]
    --quote-preset      NAME    scripture | clinical     [scripture]

Page geometry variables:
    --page-width        IN      [6]
    --page-height       IN      [9]
    --margin-inner      IN      Inner (spine-side) margin.  [0.75]
    --margin-outer      IN      Outer margin.               [0.5]
    --margin-top        IN      [0.75]
    --margin-bottom     IN      [0.75]

Misc:
    --samples           Include style sample paragraphs in the output.

Quote presets
-------------
scripture (default)
    Block Text:       Left rule only, no shading. Open and uncluttered.
    Scripture Quote:  Same base, muted colour. Applied via fenced div.

clinical
    Block Text:       Light grey shading + left rule. Higher visual weight.
    Scripture Quote:  Inherits shading; adjusted indent and spacing.
    Clinical Example: Blue-grey shading + blue rule. For case vignettes.

Quote font size is always body_size - 0.5pt regardless of preset.
"""

from __future__ import annotations

import argparse
import io
import subprocess
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.styles.style import StyleFactory

# ---------------------------------------------------------------------------
# Fixed palette — print-safe greys, not book-configurable
# ---------------------------------------------------------------------------

_DARK  = RGBColor(34,  34,  34)
_MUTED = RGBColor(88,  96,  105)
_WHITE = RGBColor(255, 255, 255)


# ---------------------------------------------------------------------------
# Book configuration
# ---------------------------------------------------------------------------

@dataclass
class BookConfig:
    """All typographic settings for one book, sourced from book.env."""
    body_font:          str   = "Palatino Linotype"
    body_size:          float = 11.0
    body_align:         str   = "justify"
    first_line_indent:  float = 0.25
    line_spacing:       float = 1.0
    h1_size:            float = 16.0
    h2_size:            float = 13.0
    h3_size:            float = 12.0
    footnote_size:      float = 9.0
    table_body_font:    str   = "Georgia"
    table_body_size:    float = 9.9
    table_header_font:  str   = "Arial"
    table_header_size:  float = 10.0
    quote_preset:       str   = "scripture"
    page_width:         float = 6.0
    page_height:        float = 9.0
    margin_inner:       float = 0.75
    margin_outer:       float = 0.5
    margin_top:         float = 0.75
    margin_bottom:      float = 0.75


def config_from_args(args: argparse.Namespace) -> BookConfig:
    return BookConfig(
        body_font         = args.body_font,
        body_size         = args.body_size,
        body_align        = args.body_align,
        first_line_indent = args.first_line_indent,
        line_spacing      = args.line_spacing,
        h1_size           = args.h1_size,
        h2_size           = args.h2_size,
        h3_size           = args.h3_size,
        footnote_size     = args.footnote_size,
        table_body_font   = args.table_body_font,
        table_body_size   = args.table_body_size,
        table_header_font = args.table_header_font,
        table_header_size = args.table_header_size,
        quote_preset      = args.quote_preset,
        page_width        = args.page_width,
        page_height       = args.page_height,
        margin_inner      = args.margin_inner,
        margin_outer      = args.margin_outer,
        margin_top        = args.margin_top,
        margin_bottom     = args.margin_bottom,
    )


# ---------------------------------------------------------------------------
# Style lookup — bypasses python-docx BabelFish name translation
# ---------------------------------------------------------------------------

def style_by_id(doc: Document, style_id: str):
    el = doc.styles._element.get_by_id(style_id)
    if el is None:
        el = doc.styles._element.get_by_id(style_id.lower())
    if el is None:
        raise KeyError(f"No style with styleId {style_id!r}")
    return StyleFactory(el)


def ensure_style(doc: Document, name: str, style_type: WD_STYLE_TYPE,
                 base_id: str | None = None):
    style_id = name.replace(" ", "")
    el = doc.styles._element.get_by_id(style_id)
    if el is not None:
        style = StyleFactory(el)
    else:
        style = doc.styles.add_style(name, style_type)
    clear_style_formatting(style)
    if base_id is not None:
        style.base_style = style_by_id(doc, base_id)
    return style


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def get_or_add_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def remove_child(parent, tag: str) -> None:
    child = parent.find(qn(tag))
    if child is not None:
        parent.remove(child)


def clear_style_formatting(style) -> None:
    style_el = style.element
    for tag in (
        "w:pPr", "w:rPr", "w:basedOn", "w:next", "w:link",
        "w:uiPriority", "w:qFormat", "w:semiHidden", "w:unhideWhenUsed",
        "w:hidden", "w:autoRedefine", "w:rsid", "w:aliases",
    ):
        remove_child(style_el, tag)


def set_style_ui(style, *, hidden=None, quick=None, priority=None) -> None:
    style.hidden      = hidden
    style.quick_style = quick
    style.priority    = priority


def set_next(doc: Document, style, next_style_id: str) -> None:
    style.next_paragraph_style = style_by_id(doc, next_style_id)


def set_outline_level(style, level: int | None) -> None:
    pPr = get_or_add_child(style.element, "w:pPr")
    remove_child(pPr, "w:outlineLvl")
    if level is not None:
        outline = OxmlElement("w:outlineLvl")
        outline.set(qn("w:val"), str(level))
        pPr.append(outline)


def set_font(style, *, name: str | None = None, size_pt: float | None = None,
             bold: bool | None = None, italic: bool | None = None,
             all_caps: bool | None = None, superscript: bool | None = None,
             color: RGBColor | None = None) -> None:
    """Set character formatting at both python-docx and raw XML layers so
    theme-font overrides cannot clobber explicit font assignments."""
    font           = style.font
    font.name      = name
    font.size      = Pt(size_pt) if size_pt is not None else None
    font.bold      = bold
    font.italic    = italic
    font.all_caps  = all_caps
    if color is not None:
        font.color.rgb = color

    rPr = get_or_add_child(style.element, "w:rPr")
    if name is not None:
        rFonts = get_or_add_child(rPr, "w:rFonts")
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            rFonts.set(qn(f"w:{attr}"), name)
        for attr in ("asciiTheme", "hAnsiTheme", "cstheme", "eastAsiaTheme"):
            rFonts.attrib.pop(qn(f"w:{attr}"), None)
    else:
        remove_child(rPr, "w:rFonts")
    if size_pt is not None:
        sz = get_or_add_child(rPr, "w:sz")
        sz.set(qn("w:val"), str(int(round(size_pt * 2))))
    else:
        remove_child(rPr, "w:sz")
    for tag, flag in (("w:b", bold), ("w:i", italic)):
        if flag is True:
            get_or_add_child(rPr, tag)
        elif flag is False:
            remove_child(rPr, tag)
    if superscript is True:
        vert = get_or_add_child(rPr, "w:vertAlign")
        vert.set(qn("w:val"), "superscript")
    elif superscript is False:
        remove_child(rPr, "w:vertAlign")


def apply_shading(style, fill: str | None) -> None:
    pPr = get_or_add_child(style.element, "w:pPr")
    shd = pPr.find(qn("w:shd"))
    if fill is None:
        if shd is not None:
            pPr.remove(shd)
        return
    if shd is None:
        shd = OxmlElement("w:shd")
        pPr.append(shd)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)


def apply_left_border(style, color: str | None, size: int) -> None:
    pPr  = get_or_add_child(style.element, "w:pPr")
    pBdr = pPr.find(qn("w:pBdr"))
    if color is None:
        if pBdr is not None:
            pPr.remove(pBdr)
        return
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    left = pBdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        pBdr.append(left)
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(size))
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), color)


def clear_borders_and_shading(style) -> None:
    pPr = get_or_add_child(style.element, "w:pPr")
    for tag in ("w:pBdr", "w:shd"):
        node = pPr.find(qn(tag))
        if node is not None:
            pPr.remove(node)


# ---------------------------------------------------------------------------
# Document geometry
# ---------------------------------------------------------------------------

def configure_document(doc: Document, cfg: BookConfig) -> None:
    section               = doc.sections[0]
    section.page_width    = Inches(cfg.page_width)
    section.page_height   = Inches(cfg.page_height)
    section.top_margin    = Inches(cfg.margin_top)
    section.bottom_margin = Inches(cfg.margin_bottom)
    section.left_margin   = Inches(cfg.margin_inner)
    section.right_margin  = Inches(cfg.margin_outer)
    section.header_distance = Inches(0.375)
    section.footer_distance = Inches(0.375)
    section.gutter          = Inches(0)
    section.start_type      = WD_SECTION_START.NEW_PAGE
    section.different_first_page_header_footer = True
    doc.settings.odd_and_even_pages_header_footer = True

    settings = doc.settings.element
    compat   = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    for cname, cval in (
        ("compatibilityMode",                          "14"),
        ("overrideTableStyleFontSizeAndJustification", "1"),
        ("enableOpenTypeFeatures",                     "1"),
        ("doNotFlipMirrorIndents",                     "1"),
    ):
        found = next(
            (el for el in compat.findall(qn("w:compatSetting"))
             if el.get(qn("w:name")) == cname), None)
        if found is None:
            found = OxmlElement("w:compatSetting")
            compat.append(found)
        found.set(qn("w:name"), cname)
        found.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
        found.set(qn("w:val"), cval)


# ---------------------------------------------------------------------------
# Normal / body text
# ---------------------------------------------------------------------------

def configure_normal(doc: Document, cfg: BookConfig) -> None:
    s = style_by_id(doc, "Normal")
    clear_style_formatting(s)
    set_style_ui(s, hidden=False, quick=True, priority=None)
    set_font(s, name=cfg.body_font, size_pt=cfg.body_size)

    align = (WD_ALIGN_PARAGRAPH.JUSTIFY
             if cfg.body_align == "justify"
             else WD_ALIGN_PARAGRAPH.LEFT)
    pf = s.paragraph_format
    pf.alignment         = align
    pf.left_indent       = None
    pf.right_indent      = None
    pf.first_line_indent = Inches(cfg.first_line_indent)
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(5)
    pf.keep_together     = None
    pf.keep_with_next    = None
    pf.page_break_before = None

    # line_spacing: 1.0 → SINGLE rule; anything else → MULTIPLE
    if cfg.line_spacing == 1.0:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.line_spacing      = 1.0
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing      = cfg.line_spacing


# ---------------------------------------------------------------------------
# Header / footer styles
# ---------------------------------------------------------------------------

def configure_header_footer_styles(doc: Document, cfg: BookConfig) -> None:
    # Running heads are always 2pt smaller than body so they read as
    # subordinate — the relationship is relative, not absolute.
    header_size = max(cfg.body_size - 2.0, 7.0)
    for name in ("Header", "Footer"):
        s = ensure_style(doc, name, WD_STYLE_TYPE.PARAGRAPH, base_id="Normal")
        set_style_ui(s, hidden=None, quick=False, priority=99)
        set_font(s, name=cfg.body_font, size_pt=header_size, color=_MUTED)
        pf = s.paragraph_format
        pf.alignment         = WD_ALIGN_PARAGRAPH.LEFT
        pf.left_indent       = None
        pf.right_indent      = None
        pf.first_line_indent = None
        pf.space_before      = None
        pf.space_after       = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.line_spacing      = 1.0


# ---------------------------------------------------------------------------
# Headings
# ---------------------------------------------------------------------------

def configure_headings(doc: Document, cfg: BookConfig) -> None:
    definitions = {
        "Heading1": dict(size=cfg.h1_size, align=WD_ALIGN_PARAGRAPH.CENTER,
                         before=24, after=12),
        "Heading2": dict(size=cfg.h2_size, align=WD_ALIGN_PARAGRAPH.LEFT,
                         before=18, after=10),
        "Heading3": dict(size=cfg.h3_size, align=WD_ALIGN_PARAGRAPH.LEFT,
                         before=14, after=8),
    }
    for style_id, d in definitions.items():
        s = style_by_id(doc, style_id)
        clear_style_formatting(s)
        s.base_style = style_by_id(doc, "Normal")
        set_next(doc, s, "BodyText")
        set_style_ui(s, hidden=False, quick=True, priority=9)
        set_font(s, name=cfg.body_font, size_pt=d["size"], bold=True)
        pf = s.paragraph_format
        pf.alignment         = d["align"]
        pf.left_indent       = None
        pf.right_indent      = None
        pf.first_line_indent = Inches(0)
        pf.space_before      = Pt(d["before"])
        pf.space_after       = Pt(d["after"])
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.line_spacing      = 1.0
        pf.keep_together     = True
        pf.keep_with_next    = True
        pf.page_break_before = None


# ---------------------------------------------------------------------------
# Quote styles — preset-driven
# ---------------------------------------------------------------------------

def _make_quote_style(doc: Document, name: str, base_id: str, cfg: BookConfig,
                      italic: bool,
                      color: RGBColor,
                      shading: str | None,
                      border_color: str | None,
                      border_size: int,
                      left_indent: float,
                      right_indent: float,
                      space_before: float,
                      space_after: float) -> None:
    """Create or update one quote paragraph style."""
    style_id = name.replace(" ", "")
    s = ensure_style(doc, style_id, WD_STYLE_TYPE.PARAGRAPH, base_id=base_id)
    s.name = name  # preserve display name with spaces

    # Quote body size tracks body_size so relative type scale is preserved.
    quote_size = max(cfg.body_size - 0.5, 7.0)
    set_font(s, name=cfg.body_font, size_pt=quote_size, italic=italic, color=color)

    pf = s.paragraph_format
    pf.left_indent       = Inches(left_indent)
    pf.right_indent      = Inches(right_indent)
    pf.first_line_indent = Inches(0)
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = 1.08

    apply_shading(s, shading)
    apply_left_border(s, border_color, border_size)


def configure_quotes(doc: Document, cfg: BookConfig) -> None:
    """
    Configure Block Text and named quote styles.

    Pandoc maps bare > blockquotes to "Block Text" in DOCX — that name is
    built into pandoc's DOCX writer and cannot be changed via reference.docx.
    Both presets style Block Text; they differ in whether shading is applied.

    All named styles (Scripture Quote, Clinical Example) are applied in
    Markdown via fenced divs:

        ::: {custom-style="Scripture Quote"}
        > For God so loved the world… (John 3:16)
        :::
    """
    if cfg.quote_preset == "clinical":
        _make_quote_style(
            doc, "Block Text", "Normal", cfg,
            italic       = True,
            color        = _DARK,
            shading      = "F3F3F3",
            border_color = "B8B8B8",
            border_size  = 14,
            left_indent  = 0.35,
            right_indent = 0.25,
            space_before = 8,
            space_after  = 12,
        )
        _make_quote_style(
            doc, "Scripture Quote", "BlockText", cfg,
            italic       = True,
            color        = _MUTED,
            shading      = "F3F3F3",
            border_color = "B8B8B8",
            border_size  = 14,
            left_indent  = 0.42,
            right_indent = 0.18,
            space_before = 12,
            space_after  = 14,
        )
        _make_quote_style(
            doc, "Clinical Example", "BlockText", cfg,
            italic       = False,
            color        = _DARK,
            shading      = "EAF0FB",    # blue-grey
            border_color = "7090C0",    # muted blue rule
            border_size  = 14,
            left_indent  = 0.35,
            right_indent = 0.25,
            space_before = 10,
            space_after  = 14,
        )
    else:
        # scripture (default) — clean, left rule only, no shading
        _make_quote_style(
            doc, "Block Text", "Normal", cfg,
            italic       = True,
            color        = _DARK,
            shading      = None,
            border_color = "C8C8C8",
            border_size  = 10,
            left_indent  = 0.35,
            right_indent = 0.25,
            space_before = 8,
            space_after  = 12,
        )
        _make_quote_style(
            doc, "Scripture Quote", "BlockText", cfg,
            italic       = True,
            color        = _MUTED,
            shading      = None,
            border_color = "B0B0B8",
            border_size  = 10,
            left_indent  = 0.42,
            right_indent = 0.18,
            space_before = 12,
            space_after  = 14,
        )


# ---------------------------------------------------------------------------
# Remaining paragraph styles
# ---------------------------------------------------------------------------

def configure_paragraph_styles(doc: Document, cfg: BookConfig) -> None:
    # Body Text
    s = ensure_style(doc, "BodyText", WD_STYLE_TYPE.PARAGRAPH, base_id="Normal")
    set_style_ui(s, hidden=None, quick=False, priority=99)
    s.paragraph_format.space_after = Pt(6)

    # First Paragraph — no first-line indent
    s = ensure_style(doc, "FirstParagraph", WD_STYLE_TYPE.PARAGRAPH, base_id="Normal")
    s.paragraph_format.first_line_indent = Inches(0)

    # Compact — no spacing, no indent
    s = ensure_style(doc, "Compact", WD_STYLE_TYPE.PARAGRAPH, base_id="BodyText")
    set_next(doc, s, "Compact")
    s.paragraph_format.first_line_indent = Inches(0)
    s.paragraph_format.space_before      = Pt(0)
    s.paragraph_format.space_after       = Pt(0)

    # Captioned Figure
    s = ensure_style(doc, "CaptionedFigure", WD_STYLE_TYPE.PARAGRAPH, base_id="Normal")
    pf = s.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Inches(0)
    pf.space_before      = Pt(6)
    pf.space_after       = Pt(3)

    # Image Caption — 1pt smaller than body, italic
    s = ensure_style(doc, "ImageCaption", WD_STYLE_TYPE.PARAGRAPH, base_id="Normal")
    set_font(s, name=cfg.body_font, size_pt=max(cfg.body_size - 1.0, 7.0), italic=True)
    pf = s.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Inches(0)
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(6)

    # TOC Heading
    s = ensure_style(doc, "TOCHeading", WD_STYLE_TYPE.PARAGRAPH, base_id="Normal")
    set_next(doc, s, "Normal")
    set_style_ui(s, hidden=True, quick=True, priority=39)
    set_outline_level(s, 9)
    set_font(s, name=cfg.body_font, size_pt=cfg.h1_size, bold=True)
    pf = s.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Inches(0)
    pf.space_before      = Pt(24)
    pf.space_after       = Pt(12)
    pf.keep_together     = True
    pf.keep_with_next    = True

    # Footnote Text
    s = ensure_style(doc, "FootnoteText", WD_STYLE_TYPE.PARAGRAPH, base_id="Normal")
    set_next(doc, s, "FootnoteText")
    set_style_ui(s, hidden=True, quick=False, priority=99)
    set_font(s, name=cfg.body_font, size_pt=cfg.footnote_size)
    pf = s.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Inches(0)
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing      = 1.0


# ---------------------------------------------------------------------------
# Lists
# ---------------------------------------------------------------------------

def configure_lists(doc: Document, cfg: BookConfig) -> None:
    # Use ensure_style rather than style_by_id: pandoc stores "List Paragraph"
    # under different styleIds across versions ("ListParagraph" vs the spaced
    # form). ensure_style strips spaces to form the lookup key, finds the
    # existing style under either convention, and creates it if absent.
    s = ensure_style(doc, "List Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    set_font(s, name=cfg.body_font, size_pt=cfg.body_size, color=_DARK)
    pf = s.paragraph_format
    pf.space_before      = Pt(2)
    pf.space_after       = Pt(5)
    pf.left_indent       = Inches(0.25)
    pf.first_line_indent = Inches(-0.18)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = 1.08


# ---------------------------------------------------------------------------
# Table styles
# ---------------------------------------------------------------------------

def configure_table_styles(doc: Document, cfg: BookConfig) -> None:
    # Table Text
    s = ensure_style(doc, "TableText", WD_STYLE_TYPE.PARAGRAPH, base_id="Normal")
    s.name = "Table Text"
    clear_borders_and_shading(s)
    set_font(s, name=cfg.table_body_font, size_pt=cfg.table_body_size, color=_DARK)
    pf = s.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.LEFT
    pf.left_indent       = Inches(0)
    pf.right_indent      = Inches(0)
    pf.first_line_indent = Inches(0)
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(1)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing      = 1.0
    pf.keep_together     = True
    pf.keep_with_next    = False

    # Table Header Text
    s = ensure_style(doc, "TableHeaderText", WD_STYLE_TYPE.PARAGRAPH,
                     base_id="TableText")
    s.name = "Table Header Text"
    clear_borders_and_shading(s)
    set_font(s, name=cfg.table_header_font, size_pt=cfg.table_header_size,
             bold=True, color=_WHITE)
    pf = s.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(1)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing      = 1.0
    pf.keep_together     = True


# ---------------------------------------------------------------------------
# Character styles
# ---------------------------------------------------------------------------

def configure_character_styles(doc: Document, cfg: BookConfig) -> None:
    # ChapterTitleOnly — STYLEREF marker for running headers
    ensure_style(doc, "ChapterTitleOnly", WD_STYLE_TYPE.CHARACTER)

    # Footnote Text Char
    s = ensure_style(doc, "FootnoteTextChar", WD_STYLE_TYPE.CHARACTER)
    set_font(s, name=cfg.body_font, size_pt=cfg.footnote_size)

    # Footnote Reference
    s = ensure_style(doc, "FootnoteReference", WD_STYLE_TYPE.CHARACTER)
    set_font(s, name=cfg.body_font, size_pt=cfg.footnote_size, superscript=True)


# ---------------------------------------------------------------------------
# Header / footer XML parts
# ---------------------------------------------------------------------------

def configure_header_footer_parts(doc: Document) -> None:
    """Initialise all six header/footer XML parts.  Content is injected later
    by postprocess-pandoc.py after the pandoc build."""
    section = doc.sections[0]
    for story, style_id in (
        (section.header,             "Header"),
        (section.first_page_header,  "Header"),
        (section.even_page_header,   "Header"),
        (section.footer,             "Footer"),
        (section.first_page_footer,  "Footer"),
        (section.even_page_footer,   "Footer"),
    ):
        element = story._element
        for child in list(element):
            element.remove(child)
        p = story.add_paragraph()
        p.style = style_by_id(doc, style_id)


# ---------------------------------------------------------------------------
# Assembly
# ---------------------------------------------------------------------------

def add_style_samples(doc: Document, cfg: BookConfig) -> None:
    samples = [
        ("Heading 1",
         f"1. Chapter Title  ({cfg.body_font} {cfg.h1_size}pt bold)"),
        ("Heading 2",
         f"Section Heading  ({cfg.h2_size}pt bold)"),
        ("Heading 3",
         f"Subsection Heading  ({cfg.h3_size}pt bold)"),
        ("Body Text",
         f"Body text  ({cfg.body_font} {cfg.body_size}pt, "
         f"indent {cfg.first_line_indent}in, spacing {cfg.line_spacing})."),
        ("First Paragraph",
         "First paragraph after heading — no first-line indent."),
        ("Block Text",
         f"Block Text: bare > blockquote  (preset: {cfg.quote_preset})."),
        ("Scripture Quote",
         "Scripture Quote: fenced div.  "
         "For God so loved the world… (John 3:16)"),
        ("List Paragraph",
         "• List item sample"),
        ("Table Text",
         f"Table body  ({cfg.table_body_font} {cfg.table_body_size}pt)"),
        ("Table Header Text",
         f"TABLE HEADER  ({cfg.table_header_font} "
         f"{cfg.table_header_size}pt bold white)"),
        ("Compact",
         "Compact paragraph — no spacing above or below."),
        ("Footnote Text",
         f"Footnote text  ({cfg.body_font} {cfg.footnote_size}pt)."),
    ]
    if cfg.quote_preset == "clinical":
        samples.insert(7, (
            "Clinical Example",
            "Clinical Example: blue-grey shaded vignette block.",
        ))
    for style_name, text in samples:
        try:
            doc.add_paragraph(text, style=style_name)
        except KeyError:
            doc.add_paragraph(f"[missing style: {style_name}] {text}")


def strip_body_for_reference(doc: Document) -> None:
    body   = doc._body._element
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        body.remove(child)
    if sectPr is not None:
        body.append(sectPr)


def build_reference_docx(output_path: str | Path,
                          cfg: BookConfig,
                          include_samples: bool = False) -> Path:
    """Build and save the reference template for a specific book configuration."""
    output_path = Path(output_path)

    # Seed from pandoc's own default reference.docx so built-in style IDs
    # (Heading1, FootnoteText, ListParagraph, etc.) are present and wired.
    result = subprocess.run(
        ["pandoc", "--print-default-data-file", "reference.docx"],
        capture_output=True, check=True,
    )
    doc = Document(io.BytesIO(result.stdout))

    configure_document(doc, cfg)
    configure_normal(doc, cfg)
    configure_header_footer_styles(doc, cfg)
    configure_headings(doc, cfg)
    configure_quotes(doc, cfg)           # Block Text and named quote styles
    configure_paragraph_styles(doc, cfg) # remaining pandoc built-in styles
    configure_lists(doc, cfg)
    configure_table_styles(doc, cfg)
    configure_character_styles(doc, cfg)
    configure_header_footer_parts(doc)

    if include_samples:
        add_style_samples(doc, cfg)

    strip_body_for_reference(doc)
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description="Build pandoc reference-template.docx from book.env settings.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    p.add_argument("output", nargs="?", default="reference-template.docx",
                   help="Output path for the .docx file")

    g = p.add_argument_group("typography")
    g.add_argument("--body-font",         default="Palatino Linotype",
                   metavar="FONT")
    g.add_argument("--body-size",         type=float, default=11.0,
                   metavar="PT")
    g.add_argument("--body-align",        default="justify",
                   choices=["justify", "left"])
    g.add_argument("--first-line-indent", type=float, default=0.25,
                   metavar="IN")
    g.add_argument("--line-spacing",      type=float, default=1.0,
                   metavar="MULT",
                   help="1.0 = single; 1.15 = 15%% looser; etc.")
    g.add_argument("--h1-size",           type=float, default=16.0, metavar="PT")
    g.add_argument("--h2-size",           type=float, default=13.0, metavar="PT")
    g.add_argument("--h3-size",           type=float, default=12.0, metavar="PT")
    g.add_argument("--footnote-size",     type=float, default=9.0,  metavar="PT")
    g.add_argument("--table-body-font",   default="Georgia",          metavar="FONT")
    g.add_argument("--table-body-size",   type=float, default=9.9,   metavar="PT")
    g.add_argument("--table-header-font", default="Arial",            metavar="FONT")
    g.add_argument("--table-header-size", type=float, default=10.0,  metavar="PT")
    g.add_argument("--quote-preset",      default="scripture",
                   choices=["scripture", "clinical"])

    g = p.add_argument_group("page geometry")
    g.add_argument("--page-width",    type=float, default=6.0,  metavar="IN")
    g.add_argument("--page-height",   type=float, default=9.0,  metavar="IN")
    g.add_argument("--margin-inner",  type=float, default=0.75, metavar="IN",
                   help="Inner (spine-side) margin")
    g.add_argument("--margin-outer",  type=float, default=0.5,  metavar="IN")
    g.add_argument("--margin-top",    type=float, default=0.75, metavar="IN")
    g.add_argument("--margin-bottom", type=float, default=0.75, metavar="IN")

    p.add_argument("--samples", action="store_true",
                   help="Populate body with style samples for review in Word")

    return p.parse_args()


if __name__ == "__main__":
    args = parse_args()
    cfg  = config_from_args(args)
    path = build_reference_docx(args.output, cfg, include_samples=args.samples)
    print(f"Created {path.resolve()}")
    print(f"  Body:     {cfg.body_font} {cfg.body_size}pt  "
          f"({'justified' if cfg.body_align == 'justify' else 'left-aligned'})  "
          f"indent={cfg.first_line_indent}in  spacing={cfg.line_spacing}")
    print(f"  Headings: H1={cfg.h1_size}pt  H2={cfg.h2_size}pt  H3={cfg.h3_size}pt")
    print(f"  Footnote: {cfg.footnote_size}pt")
    print(f"  Tables:   body={cfg.table_body_font} {cfg.table_body_size}pt  "
          f"header={cfg.table_header_font} {cfg.table_header_size}pt")
    print(f"  Quotes:   preset={cfg.quote_preset}  "
          f"size={max(cfg.body_size - 0.5, 7.0)}pt")
    print(f"  Page:     {cfg.page_width}×{cfg.page_height}in  "
          f"inner={cfg.margin_inner}  outer={cfg.margin_outer}")
